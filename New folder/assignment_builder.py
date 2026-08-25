import sqlite3
import json
import urllib.request
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# 1. دالة صياغة الشرح البشري


def generate_humanized_explanation(sql_query, raw_result):
    url = "http://localhost:11434/api/generate"
    prompt = f"""
    You are a student explaining your database assignment answer.
    Write a short, natural explanation (2 sentences) for this SQL query result.
    Do not use AI cliché words like 'Furthermore', 'Delve', 'Testament'.
    
    Query: {sql_query}
    Data: {raw_result}
    """
    payload = {"model": "qwen2.5:1.5b", "prompt": prompt, "stream": False}
    req = urllib.request.Request(url, data=json.dumps(payload).encode(
        'utf-8'), headers={'Content-Type': 'application/json'})
    try:
        with urllib.request.urlopen(req) as response:
            return json.loads(response.read().decode('utf-8'))['response'].strip()
    except Exception as e:
        return f"Explanation error: {e}"

# 2. دالة تحويل جدول النتائج إلى صورة (Snapshot)


def generate_snapshot_image(headers, data, img_path):
    fig, ax = plt.subplots(figsize=(6, max(1.2, len(data) * 0.4)))
    ax.axis('tight')
    ax.axis('off')

    table_data = [headers] + [[str(item) for item in row] for row in data]
    table = ax.table(cellText=table_data, loc='center', cellLoc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1.2, 1.4)

    # تنسيق رؤوس الجدول
    for col in range(len(headers)):
        table[(0, col)].set_facecolor('#1F4E79')
        table[(0, col)].get_text().set_color('white')
        table[(0, col)].get_text().set_weight('bold')

    plt.savefig(img_path, bbox_inches='tight', dpi=200)
    plt.close()

# 3. إعداد قاعدة بيانات Apex Hub


def setup_database():
    db_path = os.path.join(SCRIPT_DIR, "apex_media_hub.db")
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # جدول الأعضاء
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS Members (
        member_id INTEGER PRIMARY KEY AUTOINCREMENT,
        first_name TEXT NOT NULL,
        last_name TEXT NOT NULL,
        total_fines REAL DEFAULT 0.0
    );
    """)

    cursor.execute("DELETE FROM Members;")
    sample_members = [
        ('Alex', 'Tan', 180.50),
        ('Siti', 'Aishah', 45.00),
        ('John', 'Doe', 210.00),
        ('Sarah', 'Lee', 120.00),
        ('Michael', 'Wong', 310.20)
    ]
    cursor.executemany(
        "INSERT INTO Members (first_name, last_name, total_fines) VALUES (?, ?, ?);", sample_members)
    conn.commit()
    return conn

# 4. تشغيل وتوليد التقرير بالصور


def run_assignment_generator():
    conn = setup_database()
    cursor = conn.cursor()

    # Question 1: Members with Total Fines > RM150
    q1_sql = "SELECT member_id, first_name, last_name, total_fines FROM Members WHERE total_fines > 150;"
    cursor.execute(q1_sql)
    rows = cursor.fetchall()
    headers = ['Member ID', 'First Name', 'Last Name', 'Total Fines (RM)']
    conn.close()

    # حفظ صورة إثبات التشغيل
    img_filename = os.path.join(SCRIPT_DIR, "q1_snapshot.png")
    generate_snapshot_image(headers, rows, img_filename)
    print(" تم توليد صورة لقطة الشاشة (Execution Snapshot).")

    # صياغة الشرح
    explanation = generate_humanized_explanation(q1_sql, str(rows))

    # إنشاء ملف Word
    doc = Document()
    doc.add_heading('Apex Event & Media Hub - Database Assignment', level=1)
    doc.add_heading('Part 2: SQL Data Manipulation Language (DML)', level=2)

    doc.add_heading('Question 1: Members with Total Fines > RM150', level=3)

    doc.add_paragraph('SQL Command:').bold = True
    p_code = doc.add_paragraph(q1_sql)
    p_code.style = 'Quote'

    doc.add_paragraph('Execution Proof (Screenshot):').bold = True
    doc.add_picture(img_filename, width=Inches(5.0))

    doc.add_paragraph('\nExplanation of Result:').bold = True
    doc.add_paragraph(explanation)

    doc_path = os.path.join(SCRIPT_DIR, "Apex_Assignment_Part2_Output.docx")
    doc.save(doc_path)
    print(f" تم تحديث المستند بنجاح بالصور والشرح: {doc_path}")


if __name__ == "__main__":
    run_assignment_generator()
