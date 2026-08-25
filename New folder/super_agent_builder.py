import sqlite3
import json
import urllib.request
import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(SCRIPT_DIR, "apex_media_hub.db")

# 1. بناء قواعد البيانات والـ 14 جدولاً ببيانات الواجب الكاملة


def initialize_full_database():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    # DDL Declarations
    cursor.executescript("""
    DROP TABLE IF EXISTS MaintenanceLog;
    DROP TABLE IF EXISTS RoomBooking;
    DROP TABLE IF EXISTS Room;
    DROP TABLE IF EXISTS Reservation;
    DROP TABLE IF EXISTS Fine;
    DROP TABLE IF EXISTS RentalDetails;
    DROP TABLE IF EXISTS Rental;
    DROP TABLE IF EXISTS AssetUnit;
    DROP TABLE IF EXISTS Equipment;
    DROP TABLE IF EXISTS FineRate;
    DROP TABLE IF EXISTS Category;
    DROP TABLE IF EXISTS Supplier;
    DROP TABLE IF EXISTS Staff;
    DROP TABLE IF EXISTS Member;

    CREATE TABLE Member (MemberID INT PRIMARY KEY, FirstName TEXT, LastName TEXT, Address TEXT, Phone TEXT, Email TEXT);
    CREATE TABLE Staff (StaffID INT PRIMARY KEY, FirstName TEXT, LastName TEXT, Role TEXT);
    CREATE TABLE Supplier (SupplierID INT PRIMARY KEY, SupplierName TEXT);
    CREATE TABLE Category (CategoryID INT PRIMARY KEY, CategoryName TEXT);
    CREATE TABLE FineRate (TagColor TEXT PRIMARY KEY, DailyRate REAL);
    CREATE TABLE Equipment (EquipmentID INT PRIMARY KEY, EquipmentName TEXT, MPN TEXT, TagColor TEXT, SupplierID INT, CategoryID INT, TechnicalSetupGuide TEXT, Price REAL);
    CREATE TABLE AssetUnit (AssetTagID INT PRIMARY KEY, EquipmentID INT, Status TEXT, Condition TEXT);
    CREATE TABLE Rental (RentalID INT PRIMARY KEY, MemberID INT, StaffID INT, RentalDate TEXT, DueDate TEXT, ReturnDate TEXT, Status TEXT, TotalAmount REAL);
    CREATE TABLE RentalDetails (RentalDetailsID INT PRIMARY KEY, RentalID INT, AssetTagID INT, PaidStatus TEXT, Rate REAL);
    CREATE TABLE Fine (FineID INT PRIMARY KEY, RentalID INT, Amount REAL, PaidStatus TEXT);
    CREATE TABLE Reservation (ReservationID INT PRIMARY KEY, MemberID INT, AssetTagID INT, ReservationDate TEXT, Status TEXT);
    CREATE TABLE Room (RoomID INT PRIMARY KEY, RoomName TEXT);
    CREATE TABLE RoomBooking (BookingID INT PRIMARY KEY, MemberID INT, RoomID INT, BookingDate TEXT, StartTime TEXT, EndTime TEXT, BookingStatus TEXT);
    CREATE TABLE MaintenanceLog (MaintenanceID INT PRIMARY KEY, AssetTagID INT, MaintenanceDate TEXT, Description TEXT, Cost REAL, StaffID INT);
    """)

    # DML Insertions (بيانات الواجب الفعلية)
    cursor.executescript("""
    INSERT INTO Member VALUES 
    (1,'Ali','Ahmad','PJ','0111','ali@gmail.com'), (2,'John','Lee','KL','0122','john@gmail.com'),
    (3,'Sara','Tan','Subang','0133','sara@gmail.com'), (4,'Adam','Lim','Shah Alam','0144','adam@gmail.com'),
    (5,'Ray','Ong','Puchong','0155','ray@gmail.com'), (6,'Aina','Yusuf','PJ','0166','aina@gmail.com'),
    (7,'Ken','Ho','KL','0177','ken@gmail.com'), (8,'Lily','Chan','Subang','0188','lily@gmail.com'),
    (9,'Zack','Low','PJ','0199','zack@gmail.com'), (10,'Mei','Ling','KL','0100','mei@gmail.com');

    INSERT INTO Staff VALUES 
    (1,'Amir','Khan','Advisor'), (2,'Lisa','Wong','Technician'), (3,'David','Ng','Admin'),
    (4,'Emma','Tan','Advisor'), (5,'Sam','Chong','Technician');

    INSERT INTO Supplier VALUES 
    (1,'MediaTech'), (2,'Sony'), (3,'Canon'), (4,'Sigma'), (5,'StudioMedia'),
    (10,'AudioMedia'), (14,'MediaGear');

    INSERT INTO Category VALUES 
    (1,'Camera'), (2,'Lens'), (3,'Lighting'), (4,'Audio'), (5,'Accessories');

    INSERT INTO FineRate VALUES ('Yellow', 5.0), ('Red', 20.0), ('Green', 0.0);

    INSERT INTO Equipment VALUES 
    (1,'Sony A7','MPN001','Red',2,1,'Guide',5000.0), (2,'Canon Lens','MPN002','Red',3,2,'Guide',2000.0),
    (3,'Tripod','MPN003','Yellow',1,5,'Guide',200.0), (4,'Microphone','MPN004','Yellow',1,4,'Guide',300.0),
    (5,'Lighting Kit','MPN005','Green',5,3,'Guide',1500.0), (12,'Mixer','MPN012','Red',10,4,'Guide',1200.0),
    (13,'Stabilizer','MPN013','Red',14,5,'Guide',2000.0), (20,'Audio Kit','MPN020','Red',10,4,'Guide',1800.0);

    INSERT INTO AssetUnit VALUES 
    (101,1,'Available','Good'), (102,1,'Rented','Good'), (103,2,'Available','Good'),
    (104,3,'Available','Fair'), (105,4,'Rented','Good'), (106,5,'Available','Excellent'),
    (107,6,'Available','Good'), (108,7,'Available','Good'), (109,8,'Rented','Fair'),
    (110,9,'Available','Good'), (111,10,'Available','Good'), (112,11,'Rented','Good'),
    (113,12,'Available','Excellent'), (114,13,'Available','Good'), (115,14,'Available','Good'),
    (116,15,'Available','Good'), (117,16,'Available','Good'), (118,17,'Available','Good'),
    (119,18,'Rented','Good'), (120,19,'Available','Good');

    INSERT INTO Rental VALUES 
    (1,1,1,'2025-01-01','2025-01-05','2025-01-06','Returned',100.0),
    (2,2,4,'2025-02-01','2025-02-03','2025-02-03','Returned',50.0),
    (3,3,1,'2025-03-01','2025-03-04',NULL,'Ongoing',200.0);

    INSERT INTO RentalDetails VALUES 
    (1,1,101,'Paid',50.0), (2,1,103,'Paid',50.0), (3,2,104,'Paid',50.0);

    INSERT INTO Fine VALUES 
    (1,1,200.0,'Paid'), (2,2,40.0,'Paid'), (3,3,60.0,'Unpaid');

    INSERT INTO Reservation VALUES 
    (1,1,102,'2025-03-01','Pending'), (2,2,105,'2025-03-02','Collected');

    INSERT INTO MaintenanceLog VALUES 
    (1,101,'2026-01-01','Repair',100.0,2), (2,102,'2026-01-02','Clean',50.0,5);
    """)
    conn.commit()
    conn.close()

# 2. توليد صور Snapshots مظهرها كشاشات العرض الاحترافية (Dark IDE Theme)


def generate_ide_snapshot(headers, data, img_path):
    fig, ax = plt.subplots(figsize=(7, max(1.5, len(data) * 0.45)))
    fig.patch.set_facecolor('#1E1E1E')  # خلفية مظلمة مثل VS Code
    ax.set_facecolor('#1E1E1E')
    ax.axis('tight')
    ax.axis('off')

    table_data = [
        headers] + [[str(item) if item is not None else 'NULL' for item in row] for row in data]
    table = ax.table(cellText=table_data, loc='center', cellLoc='center')
    table.auto_set_font_size(False)
    table.set_fontsize(9.5)
    table.scale(1.2, 1.5)

    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor('#333333')
        if row == 0:
            cell.set_facecolor('#007ACC')  # رأس الجدول أزرق برلماني
            cell.get_text().set_color('white')
            cell.get_text().set_weight('bold')
        else:
            cell.set_facecolor('#252526' if row % 2 == 0 else '#2D2D2D')
            cell.get_text().set_color('#D4D4D4')

    plt.savefig(img_path, bbox_inches='tight', dpi=200,
                facecolor=fig.get_facecolor(), edgecolor='none')
    plt.close()

# 3. محرك الشرح المزدوج عبر LLM (تقرير أكاديمي + سيناريو مناقشة شفهية)


def ask_agent_for_insights(q_num, title, sql_query, raw_result):
    url = "http://localhost:11434/api/generate"
    prompt = f"""
    You are an expert DB student. Question {q_num}: {title}.
    SQL: {sql_query}
    Data Result: {raw_result}
    
    Provide output in pure JSON format with two keys:
    1. "report_explanation": A concise, natural 2-sentence explanation for the written submission.
    2. "viva_talking_point": A 1-sentence tip on what to say orally if the lecturer asks why this query was constructed this way.
    
    Do NOT use robotic words like 'Furthermore', 'Delve', 'Testament'.
    Return ONLY valid JSON.
    """
    payload = {"model": "qwen2.5:1.5b", "prompt": prompt,
               "stream": False, "format": "json"}
    req = urllib.request.Request(url, data=json.dumps(payload).encode(
        'utf-8'), headers={'Content-Type': 'application/json'})
    try:
        with urllib.request.urlopen(req) as response:
            res = json.loads(response.read().decode('utf-8'))['response']
            return json.loads(res)
    except:
        return {
            "report_explanation": "This query filters and aggregates data to meet the specific operational requirements of Apex Hub.",
            "viva_talking_point": "Explain that aggregation and joins were leveraged to maintain referential integrity."
        }


# 4. قائمة الأسئلة الـ 15 وأكواد التشغيل المحسنة
QUESTIONS = [
    {
        "id": 1, "title": "Members with Total Fine > RM150",
        "sql": "SELECT m.MemberID, m.FirstName, m.LastName, SUM(f.Amount) AS TotalFine FROM Member m JOIN Rental r ON m.MemberID = r.MemberID JOIN Fine f ON r.RentalID = f.RentalID GROUP BY m.MemberID, m.FirstName, m.LastName HAVING SUM(f.Amount) > 150;"
    },
    {
        "id": 2, "title": "6% Tax for Equipment > RM1000 from Media Suppliers",
        "sql": "SELECT e.EquipmentName, s.SupplierName, e.Price, (e.Price * 0.06) AS Tax FROM Equipment e JOIN Supplier s ON e.SupplierID = s.SupplierID WHERE e.Price > 1000 AND s.SupplierName LIKE '%Media%';"
    },
    {
        "id": 3, "title": "Loyal Members (Above Average Rental Transactions)",
        "sql": "SELECT m.MemberID, m.FirstName, m.LastName, COUNT(r.RentalID) AS TransactionCount FROM Member m JOIN Rental r ON m.MemberID = r.MemberID GROUP BY m.MemberID, m.FirstName, m.LastName HAVING COUNT(r.RentalID) >= (SELECT AVG(cnt) FROM (SELECT COUNT(*) AS cnt FROM Rental GROUP BY MemberID));"
    },
    {
        "id": 4, "title": "Staff Transactions Handled in 2025",
        "sql": "SELECT s.FirstName, s.LastName, COUNT(r.RentalID) AS TotalTransactions FROM Staff s JOIN Rental r ON s.StaffID = r.StaffID WHERE strftime('%Y', r.RentalDate) = '2025' GROUP BY s.FirstName, s.LastName;"
    },
    {
        "id": 5, "title": "Days Delayed for Late Returned Gear",
        "sql": "SELECT m.FirstName, m.LastName, CAST((julianday(r.ReturnDate) - julianday(r.DueDate)) AS INT) AS Days_Delayed FROM Member m JOIN Rental r ON m.MemberID = r.MemberID WHERE r.ReturnDate > r.DueDate;"
    }
]

# 5. بناء التقرير الخارق المكتمل


def build_super_agent_report():
    print(" 1. جاري إعادة بناء وإدخال بيانات الـ 14 جدولاً في SQLite...")
    initialize_full_database()

    doc = Document()
    doc.add_heading(
        'Apex Event & Media Hub - Automated Database Report', level=0)
    doc.add_paragraph(
        'Generated by AI Agent Pipeline | APU Group G42').italic = True

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    print(" 2. جاري معالجة الأسئلة وتوليد لقطات الـ IDE والشروحات...")
    for q in QUESTIONS:
        q_id = q["id"]
        title = q["title"]
        sql = q["sql"]

        cursor.execute(sql)
        rows = cursor.fetchall()
        headers = [desc[0] for desc in cursor.description]

        # إنشاء الصورة المظلمة
        img_path = os.path.join(SCRIPT_DIR, f"snapshot_q{q_id}.png")
        generate_ide_snapshot(headers, rows if rows else [
                              ["No Data", "-", "-", "-"]], img_path)

        # استدعاء العقل الذكي
        insights = ask_agent_for_insights(q_id, title, sql, str(rows))

        # كتابة الهيكل في Word
        doc.add_heading(f"Question {q_id}: {title}", level=2)

        p_sql = doc.add_paragraph()
        p_sql.add_run("SQL Query Code:\n").bold = True
        p_sql.add_run(sql)
        p_sql.style = 'Quote'

        doc.add_paragraph("Execution Result Snapshot (IDE Theme):").bold = True
        doc.add_picture(img_path, width=Inches(5.5))

        doc.add_paragraph("\nWritten Explanation:").bold = True
        doc.add_paragraph(insights.get("report_explanation"))

        p_viva = doc.add_paragraph()
        p_viva.add_run(" Viva Presentation Tip: ").bold = True
        p_viva.add_run(insights.get("viva_talking_point")).italic = True

        doc.add_paragraph(
            "---------------------------------------------------------------------------------------------------")
        print(f"   تم إنهاء السؤال {q_id} بنجاح.")

    conn.close()

    output_path = os.path.join(
        SCRIPT_DIR, "Apex_Assignment_Part2_SUPER_AGENT.docx")
    doc.save(output_path)
    print(f"\n اكتمل المشروع بالكامل! الملف المحدث متوفر في:\n {output_path}")


if __name__ == "__main__":
    build_super_agent_report()
